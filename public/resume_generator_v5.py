#!/home/redemption/codebase/docs/auto-gen-scripts/venv/bin/python

"""
Resume Generator v5
- Target A: Global Frontend (FE) -> React/Next.js expert.
- Target B: Global Fullstack (FS) -> Node/Python/React expert.
- Target C: Global Backend (BE) -> NestJS/Express/Laravel/System Architecture expert.
- Target D: Mobile (Mobile) -> React Native/iOS/Android expert (Senior Framing).
- Constraint: NO Founder/Nexalith references.
"""

import argparse
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# -------------------------------------------------------------------------
# CONSTANTS & STYLES
# -------------------------------------------------------------------------
NAME = 'REDEMPTION JONATHAN'
CONTACT_INFO = 'Nigeria | redemptionjonathan1@gmail.com | +234 903 4847 432'
LINKS = 'LinkedIn: https://linkedin.com/in/redemption | GitHub: https://github.com/jrcity'

def set_base_styles(doc: Document):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri' 
    font.size = Pt(11)

def add_header(doc: Document, role_title: str):
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    name_run = header.add_run(NAME + "\n")
    name_run.bold = True
    name_run.font.size = Pt(24)
    name_run.font.color.rgb = RGBColor(0, 0, 0)

    title_run = header.add_run(role_title + "\n")
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_run.font.color.rgb = RGBColor(89, 89, 89)

    contact_run = header.add_run(CONTACT_INFO + "\n")
    contact_run.font.size = Pt(10)

    links_run = header.add_run(LINKS)
    links_run.font.size = Pt(10)
    links_run.font.color.rgb = RGBColor(0, 51, 204)

    doc.add_paragraph()

def add_heading(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.style = 'Heading 1'

def add_skills_table(doc: Document, skills_rows):
    heading = doc.add_paragraph('TECHNICAL SKILLS')
    heading.style = 'Heading 1'

    table = doc.add_table(rows=0, cols=2)
    table.autofit = False 
    table.columns[0].width = Inches(1.5)
    table.columns[1].width = Inches(5.0)

    for category, items in skills_rows:
        row_cells = table.add_row().cells
        cat_run = row_cells[0].paragraphs[0].add_run(category)
        cat_run.bold = True
        row_cells[1].text = items
    
    doc.add_paragraph()

def add_section(doc: Document, title: str, data: list):
    add_heading(doc, title)
    for item in data:
        p = doc.add_paragraph()
        
        # Title
        title_run = p.add_run(item['title'])
        title_run.bold = True
        title_run.font.size = Pt(12)
        
        # Org
        if item.get('org'):
            p.add_run(f" | {item['org']}").italic = True
        
        # Date (Right Aligned via tabs)
        if item.get('date'):
            p.add_run(f"	{item['date']}").bold = True
            p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_ALIGN_PARAGRAPH.RIGHT)
        
        # Bullets
        for bullet in item.get('bullets', []):
            b_p = doc.add_paragraph(bullet, style='List Bullet')
            b_p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()

def add_education(doc: Document):
    add_heading(doc, 'EDUCATION')
    p = doc.add_paragraph()
    p.add_run('B.Eng. Mechatronics and Systems Engineering').bold = True
    p.add_run('\nAbubakar Tafawa Balewa University, Bauchi').italic = True
    p.add_run('	Expected 2026').bold = True
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_ALIGN_PARAGRAPH.RIGHT)

# -------------------------------------------------------------------------
# DATA REPOSITORIES
# -------------------------------------------------------------------------

# --- PROJECTS ---

PROJECTS_BACKEND = [
    {
        'title': 'CashWorxs Fullstack Fintech Platform',
        'org': 'Laravel / PHP / MySQL / Blade / Tailwind',
        'bullets': [
            'Architected and single-handedly developed a comprehensive fintech ecosystem for tax and fee processing, serving as a primary payment bridge for multi-tenant organizations.',
            'Engineered a multi-gateway payment abstraction layer (PayStack, etc.) with automatic fallback logic and atomic transactional integrity for high-concurrency ledger updates.',
            'Designed a robust micro-role permission system (RBAC) and OAuth-secured API, enabling secure data sharing between State Internal Revenue Services (IRS) and third-party platforms.',
            'Built an interactive administrative dashboard and onboarding flow using Laravel Blade, Tailwind CSS, and Vite, featuring real-time financial reporting and document management.'
        ]
    },
    {
        'title': 'SnapShop Price Comparison Engine',
        'org': 'NestJS / Puppeteer / MongoDB',
        'bullets': [
            'Architected an automated scraping engine using Puppeteer to aggregate real-time product data from Jumia and Konga.',
            'Implemented optimized MongoDB text search and price comparison algorithms, reducing query latency by 50%.'
        ]
    },
    {
        'title': 'Support24 Healthcare Marketplace API',
        'org': 'Express / TypeScript / MongoDB',
        'bullets': [
            'Developed a secure RESTful API for connecting care workers with clients, implementing JWT-based RBAC for multiple user roles.',
            'Integrated Stripe for automated payment processing and AWS S3 for secure medical document storage.'
        ]
    }
]

PROJECTS_WEB = [
    {
        'title': 'AFRUNA E-commerce Platform',
        'org': 'Frontend Lead',
        'bullets': [
            'Architected dynamic product pages with Next.js 13, achieving a 98/100 Lighthouse performance score.',
            'Integrated Stripe and PayPal gateways with robust error handling, processing secure transactions.',
            'Implemented real-time search with debouncing and caching strategies to minimize API calls.'
        ]
    }
]

PROJECTS_MOBILE = [
    {
        'title': 'PoketDrive Logistics Platform',
        'org': 'Lead Mobile Architect',
        'bullets': [
            'Engineered a real-time logistics tracking app using React Native Maps and Socket.io, handling concurrent socket connections for 500+ active drivers.',
            'Implemented background location services and geofencing triggers to automate delivery status updates.'
        ]
    }
]

# --- EXPERIENCE ---

EXPERIENCE_BACKEND = [
    {
        'title': 'Senior Backend Engineer',
        'org': 'Novada Tech (Remote)',
        'date': 'June 2025 – Present',
        'bullets': [
            'Spearheading the decomposition of monolithic legacy systems into high-performance microservices using NestJS and Docker.',
            'Optimizing database performance through advanced indexing and query profiling in PostgreSQL and MongoDB, resulting in a 30% reduction in API response times.',
            'Designing and maintaining robust CI/CD pipelines with GitHub Actions, ensuring zero-downtime deployments to AWS environments.'
        ]
    },
    {
        'title': 'Backend Developer (Freelance)',
        'org': 'Global Remote',
        'date': 'Jan 2021 – Present',
        'bullets': [
            'Architected and deployed scalable backend solutions for international clients using Node.js, Python, and Laravel.',
            'Integrated complex third-party APIs (Stripe, Twilio, SendGrid) and implemented custom authentication protocols.'
        ]
    }
]

EXPERIENCE_WEB = [
    {
        'title': 'Full Stack Software Engineer',
        'org': 'Novada Tech (Remote)',
        'date': 'June 2025 – Present',
        'bullets': [
            'Spearheading the migration of legacy monolithic frontends to modular micro-frontends using React and Module Federation.',
            'Implementing extensive unit (Jest) and E2E (Cypress) testing suites, increasing code coverage to 85%.',
            'Optimizing CI/CD pipelines on GitHub Actions, reducing deployment time by 40%.'
        ]
    },
    {
        'title': 'React Developer',
        'org': 'Blares Technologies (Remote)',
        'date': 'June 2024 – Jan 2025',
        'bullets': [
            'Delivered pixel-perfect cross-platform components ensuring 100% design fidelity.',
            'Reduced app bundle size by 30% through tree-shaking and efficient asset management.'
        ]
    }
]

EXPERIENCE_MOBILE = [
    {
        'title': 'Senior Mobile Engineer (React Native)',
        'org': 'Novada Tech (Remote)',
        'date': 'June 2025 – Present',
        'bullets': [
            'Architected a scalable mobile healthcare platform using React Native, achieving 99.9% crash-free sessions.',
            'Engineered custom Native Modules (Java/Swift) to bridge legacy SDKs, improving performance by 40%.',
            'Established automated CI/CD pipelines using Fastlane, enabling one-click deployments to App Store and Play Store.'
        ]
    }
]

# -------------------------------------------------------------------------
# BUILDERS
# -------------------------------------------------------------------------

def build_backend(doc: Document):
    add_header(doc, 'SENIOR BACKEND ENGINEER')
    doc.add_paragraph(
        "Backend Specialist with 5+ years of experience building scalable distributed systems, RESTful APIs, "
        "and automated data pipelines. Expert in Node.js (NestJS/Express) and Python, with a focus on "
        "microservices architecture, system security, and cloud infrastructure optimization."
    )
    skills = [
        ("Languages", "TypeScript, JavaScript, Python, PHP, SQL, Go (Basic)"),
        ("Frameworks", "NestJS, Express.js, Laravel, FastAPI, Django"),
        ("Databases", "PostgreSQL, MongoDB (Mongoose), Redis, MySQL, DynamoDB"),
        ("Infra/DevOps", "Docker, Kubernetes, AWS (S3/EC2/Lambda), CI/CD, RabbitMQ, Microservices")
    ]
    add_skills_table(doc, skills)
    add_section(doc, 'PROFESSIONAL EXPERIENCE', EXPERIENCE_BACKEND)
    add_section(doc, 'TECHNICAL PROJECTS', PROJECTS_BACKEND)
    add_education(doc)

def build_global_fe(doc: Document):
    add_header(doc, 'SENIOR FRONTEND ENGINEER')
    doc.add_paragraph(
        "Performance-obsessed Frontend Engineer with specialized expertise in the React ecosystem (Next.js, TypeScript). "
        "Proven track record of delivering scalable web applications for remote teams. Expert in UI/UX optimization, "
        "accessibility standards (WCAG), and modern state management."
    )
    skills = [
        ("Core", "TypeScript, JavaScript (ES6+), HTML5, CSS3, JSON"),
        ("Frameworks", "React.js, Next.js (App Router), Vue.js"),
        ("Styling", "Tailwind CSS, SCSS, Styled Components, Framer Motion"),
        ("Tools/Cloud", "Git, Docker, AWS (S3/CloudFront), CI/CD, Webpack, Vite")
    ]
    add_skills_table(doc, skills)
    add_section(doc, 'PROFESSIONAL EXPERIENCE', EXPERIENCE_WEB)
    add_section(doc, 'TECHNICAL PROJECTS', PROJECTS_WEB)
    add_education(doc)

def build_global_fs(doc: Document):
    add_header(doc, 'FULL STACK SOFTWARE ENGINEER')
    doc.add_paragraph(
        "Full Stack Engineer capable of owning the entire development lifecycle. Expertise in bridging "
        "frontend interactivity (React/Next.js) with robust backend architecture (Node.js/Python). "
        "Focus on writing clean, maintainable, and testable code for scalable systems."
    )
    skills = [
        ("Languages", "TypeScript, JavaScript, Python, PHP, SQL"),
        ("Frontend", "React.js, Next.js, Redux, Tailwind CSS"),
        ("Backend", "Node.js (Express/NestJS), Laravel, REST APIs, GraphQL"),
        ("Data/Infra", "PostgreSQL, MongoDB, Redis, Docker, AWS, Firebase")
    ]
    add_skills_table(doc, skills)
    add_section(doc, 'PROFESSIONAL EXPERIENCE', EXPERIENCE_WEB)
    add_section(doc, 'TECHNICAL PROJECTS', PROJECTS_BACKEND[:2] + PROJECTS_WEB)
    add_education(doc)

def build_mobile(doc: Document):
    add_header(doc, 'SENIOR REACT NATIVE ENGINEER')
    doc.add_paragraph(
        "Senior Mobile Engineer with 6+ years of experience architecting cross-platform applications. "
        "Expert in the React Native ecosystem, specializing in bridging native modules, performance profiling, "
        "and offline-first architectures. Proven ability to lead mobile initiatives."
    )
    skills = [
        ("Core Mobile", "React Native (CLI/Expo), TypeScript, JavaScript (ES6+)"),
        ("Native & Bridge", "Java (Android), Swift (iOS), Native Modules, Gradle/CocoaPods"),
        ("State/Data", "Redux Toolkit, Zustand, React Query, Realm, WatermelonDB"),
        ("Testing & CI/CD", "Jest, Detox, Maestro, Fastlane, GitHub Actions, CodePush")
    ]
    add_skills_table(doc, skills)
    add_section(doc, 'PROFESSIONAL EXPERIENCE', EXPERIENCE_MOBILE)
    add_section(doc, 'MOBILE ARCHITECTURE PROJECTS', PROJECTS_MOBILE)
    add_education(doc)

# -------------------------------------------------------------------------
# MAIN
# -------------------------------------------------------------------------

def create_resume(target: str) -> str:
    doc = Document()
    set_base_styles(doc)

    if target == 'fe':
        build_global_fe(doc)
        filename = 'FE_Redemption_Jonathan_Resume.docx'
    elif target == 'be':
        build_backend(doc)
        filename = 'BE_Redemption_Jonathan_Resume.docx'
    elif target == 'mobile':
        build_mobile(doc)
        filename = 'Mobile_Redemption_Jonathan_Resume.docx'
    else:
        build_global_fs(doc)
        filename = 'FS_Redemption_Jonathan_Resume.docx'

    doc.save(filename)
    return filename

if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Generate Targeted Resumes')
    group = parser.add_mutually_exclusive_group(required=True)
    group.add_argument('--fe', action='store_true', help='Target: Global Frontend')
    group.add_argument('--be', action='store_true', help='Target: Global Backend')
    group.add_argument('--fs', action='store_true', help='Target: Global Fullstack')
    group.add_argument('--mobile', action='store_true', help='Target: Senior Mobile')
    
    args = parser.parse_args()
    
    if args.fe:
        target = 'fe'
    elif args.be:
        target = 'be'
    elif args.mobile:
        target = 'mobile'
    else:
        target = 'fs'
        
    out_file = create_resume(target)
    print(f"Generated Resume for {target.upper()} target: {out_file}")
